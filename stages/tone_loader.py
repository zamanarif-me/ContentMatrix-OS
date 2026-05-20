"""
Tone Loader — extracts brand voice / writing tone from uploaded sample files.

Supports:
  - .md / .markdown / .txt   → plain read
  - .docx                    → python-docx paragraph extraction
  - .zip                     → extract all md/docx/txt inside, concatenate

The extracted text is stored as a long string and passed as
`brand_voice_notes` to the section writer / refiner prompts, so the LLM
sees real examples of the brand voice and mimics them.

Priority order in pipeline:
  Sessions-level tone override (per-session)
      ↓ falls back to
  Home-level global tone (applies to all sessions)
      ↓ falls back to
  Manual text in business voice_notes
      ↓ falls back to
  None (default house rules only)
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from typing import IO, Optional


MAX_TONE_CHARS = 50_000  # cap total tone text — keeps prompt under token limits


# ── Public API ────────────────────────────────────────────────────────────────

def extract_text_from_upload(file_obj: IO, filename: str | None = None) -> str:
    """
    Given an uploaded file (Streamlit UploadedFile or similar),
    return clean plain text suitable for use as brand voice notes.

    Dispatch by extension:
      .md / .markdown / .txt    → direct read
      .docx                     → python-docx
      .zip                      → all matching files concatenated
    """
    name = (filename or getattr(file_obj, "name", "") or "").lower()
    ext = Path(name).suffix.lower()

    try:
        if ext in (".md", ".markdown", ".txt"):
            return _read_text(file_obj)
        if ext == ".docx":
            return _read_docx(file_obj)
        if ext == ".zip":
            return _read_zip(file_obj)
        # Unknown — try as plain text
        return _read_text(file_obj)
    except Exception as exc:
        raise ValueError(f"Could not parse {name}: {exc}") from exc


def combine_tone_sources(
    sessions_text: Optional[str] = None,
    home_text:     Optional[str] = None,
    manual_text:   Optional[str] = None,
) -> Optional[str]:
    """
    Apply tone priority: Sessions > Home > Manual > None.
    Returns the FIRST non-empty source (NOT a concatenation — explicit override).

    This matches the user's requirement:
    "if anyone sets the tone in the home dashboard, & also in the Sessions
     dashboard, you will work/create with sessions tone"
    """
    for candidate in (sessions_text, home_text, manual_text):
        if candidate and candidate.strip():
            return candidate.strip()[:MAX_TONE_CHARS]
    return None


# ── Format-specific parsers ──────────────────────────────────────────────────

def _read_text(file_obj: IO) -> str:
    data = file_obj.read()
    if isinstance(data, bytes):
        # Try UTF-8 first, fall back to latin-1
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("latin-1", errors="replace")
    return str(data)


def _read_docx(file_obj: IO) -> str:
    """Extract paragraph text from a .docx using python-docx."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed. Install with: pip install python-docx"
        ) from exc

    # python-docx needs a file-like object
    if hasattr(file_obj, "read"):
        doc = Document(file_obj)
    else:
        doc = Document(io.BytesIO(file_obj))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _read_zip(file_obj: IO) -> str:
    """Concatenate all text/markdown/docx files inside a zip."""
    chunks: list[str] = []
    with zipfile.ZipFile(file_obj) as zf:
        for info in zf.infolist():
            name = info.filename.lower()
            if name.endswith("/") or info.is_dir():
                continue
            ext = Path(name).suffix.lower()
            if ext not in (".md", ".markdown", ".txt", ".docx"):
                continue
            try:
                with zf.open(info) as inner:
                    if ext == ".docx":
                        # python-docx needs a real file-like — re-wrap bytes
                        chunks.append(
                            f"\n\n# {Path(name).name}\n\n"
                            + _read_docx(io.BytesIO(inner.read()))
                        )
                    else:
                        chunks.append(
                            f"\n\n# {Path(name).name}\n\n"
                            + _read_text(inner)
                        )
            except Exception:
                continue
    return "\n".join(chunks).strip()
