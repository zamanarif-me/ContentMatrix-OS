"""
Stage 8: Exporter

Renders the final GeneratedArticle into downloadable formats:
  - Markdown  (article.md)
  - HTML      (article.html)
  - DOCX      (article.docx)
  - JSON      (article.json, with full metadata)
"""

from __future__ import annotations

import io
import json
from pathlib import Path

from content_models import ExportFormat, GeneratedArticle


# ── Assembly ──────────────────────────────────────────────────────────────────

def assemble_markdown(article: GeneratedArticle) -> str:
    """Concatenate sections into a single markdown document."""
    parts: list[str] = []
    parts.append(f"# {article.title}\n")
    if article.meta_description:
        parts.append(f"_{article.meta_description}_\n")
    for s in article.sections:
        body = s.content_md.strip()
        if body and not body.startswith("#"):
            body = f"## {s.heading.text}\n\n{body}"
        parts.append(body)
    return "\n\n".join(parts)


# ── Exporters ─────────────────────────────────────────────────────────────────

def to_markdown(article: GeneratedArticle) -> bytes:
    md = article.final_md or assemble_markdown(article)
    return md.encode("utf-8")


def to_html(article: GeneratedArticle) -> bytes:
    try:
        import markdown as md_lib
    except ImportError:
        return b"<html><body>Install 'markdown' package for HTML export.</body></html>"
    md = article.final_md or assemble_markdown(article)
    html = md_lib.markdown(md, extensions=["tables", "fenced_code"])
    full = (
        f"<!DOCTYPE html><html><head>"
        f"<meta charset='utf-8'><title>{article.title}</title>"
        f"<meta name='description' content='{article.meta_description}'>"
        f"</head><body>{html}</body></html>"
    )
    return full.encode("utf-8")


def to_docx(article: GeneratedArticle) -> bytes:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    doc.add_heading(article.title, level=1)
    if article.meta_description:
        p = doc.add_paragraph(article.meta_description)
        p.italic = True
    for s in article.sections:
        doc.add_heading(s.heading.text, level=int(s.heading.level.lstrip("H")))
        for para in s.content_md.split("\n\n"):
            if para.strip().startswith("#"):
                continue
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_json(article: GeneratedArticle) -> bytes:
    return json.dumps(article.model_dump(mode="json"), indent=2).encode("utf-8")


# ── Bundle ────────────────────────────────────────────────────────────────────

def export_bundle(
    article: GeneratedArticle,
    formats: list[ExportFormat],
    out_dir: str | Path,
) -> dict[str, Path]:
    """Write all requested formats to disk. Returns {format: path}."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    written: dict[str, Path] = {}

    handlers = {
        ExportFormat.MARKDOWN: ("article.md",   to_markdown),
        ExportFormat.HTML:     ("article.html", to_html),
        ExportFormat.DOCX:     ("article.docx", to_docx),
        ExportFormat.JSON:     ("article.json", to_json),
    }

    for fmt in formats:
        if fmt not in handlers:
            continue
        filename, fn = handlers[fmt]
        path = out_dir / filename
        path.write_bytes(fn(article))
        written[fmt.value] = path

    return written
